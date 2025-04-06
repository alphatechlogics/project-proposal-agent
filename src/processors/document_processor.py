from typing import Optional
from pathlib import Path
import docx
from openpyxl import Workbook

from ..config.settings import settings

class DocumentProcessor:
    @staticmethod
    def read_docx(file_path: Path) -> str:
        """Read content from a Word document."""
        doc = docx.Document(file_path)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    
    @staticmethod
    def write_docx(content: str, output_path: Path) -> None:
        """Write content to a Word document."""
        doc = docx.Document()
        doc.add_paragraph(content)
        doc.save(output_path)
    
    @staticmethod
    def write_excel(data: dict, output_path: Path) -> None:
        """Write data to an Excel file."""
        wb = Workbook()
        ws = wb.active
        
        # Add headers
        headers = list(data.keys())
        for col, header in enumerate(headers, 1):
            ws.cell(row=1, column=col, value=header)
        
        # Add data
        for row, values in enumerate(zip(*data.values()), 2):
            for col, value in enumerate(values, 1):
                ws.cell(row=row, column=col, value=value)
        
        wb.save(output_path)
    
    @staticmethod
    def process_input_file(file_path: Path) -> Optional[str]:
        """Process an input file and return its content."""
        if not file_path.exists():
            return None
        
        if file_path.suffix.lower() == '.docx':
            return DocumentProcessor.read_docx(file_path)
        
        # Add support for other file types here
        return None