from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from loguru import logger
import io

class DocumentGenerator:
    def __init__(self):
        self.sections = {
            'requirements': self._add_requirements_section,
            'tech_specs': self._add_technical_specs_section,
            'project_plan': self._add_project_plan_section,
            'cost_estimate': self._add_cost_estimate_section
        }

    def generate_documents(self, **sections):
        try:
            doc = Document()
            self._add_title(doc)
            self._add_toc(doc)
            
            for section_name, content in sections.items():
                if section_name in self.sections and content:
                    self.sections[section_name](doc, content)
                    
            return self._save_document(doc)
        except Exception as e:
            logger.error(f"Error generating document: {str(e)}")
            raise

    def _add_title(self, doc):
        title = doc.add_heading('Project Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph().add_run().add_break()

    def _add_toc(self, doc):
        toc = doc.add_paragraph('Table of Contents')
        toc.style = 'Heading 1'
        doc.add_paragraph().add_run().add_break()

    def _add_requirements_section(self, doc, content):
        heading = doc.add_heading('Requirements Analysis', 1)
        doc.add_paragraph(content)
        doc.add_page_break()

    def _add_technical_specs_section(self, doc, content):
        heading = doc.add_heading('Technical Specifications', 1)
        doc.add_paragraph(content)
        doc.add_page_break()

    def _add_project_plan_section(self, doc, content):
        heading = doc.add_heading('Project Plan', 1)
        if isinstance(content, dict):
            for key, value in content.items():
                subheading = doc.add_heading(key.replace('_', ' ').title(), 2)
                doc.add_paragraph(str(value))
        else:
            doc.add_paragraph(str(content))
        doc.add_page_break()

    def _add_cost_estimate_section(self, doc, content):
        heading = doc.add_heading('Cost Estimate', 1)
        if isinstance(content, dict):
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Item'
            header_cells[1].text = 'Cost'
            
            for item, cost in content.items():
                row_cells = table.add_row().cells
                row_cells[0].text = str(item)
                row_cells[1].text = str(cost)
        else:
            doc.add_paragraph(str(content))

    def _save_document(self, doc):
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output.getvalue()